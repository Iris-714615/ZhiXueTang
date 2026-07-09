"""Word 文档处理器。

基于 python-docx 解析 .docx 文件，支持段落与表格内容提取，
同样使用多线程分片提升大文档处理速度。
"""
import os
from concurrent.futures import ThreadPoolExecutor
from typing import List

from llama_index.core import Document

from app.data_processing.base import BaseDataProcessor


class WordProcessor(BaseDataProcessor):
    """Word 文档多线程分片处理器。"""

    def __init__(self, max_workers: int = 4):
        self.max_workers = max_workers

    def load_and_split(self, source: str, chunk_size: int = 512, chunk_overlap: int = 100) -> List[Document]:
        """读取 Word 文档并分片。

        Args:
            source: .docx 文件路径。
            chunk_size: 分片大小。
            chunk_overlap: 分片重叠量。

        Returns:
            List[Document]: 分片后的文档列表。
        """
        if not os.path.exists(source):
            raise FileNotFoundError(f"Word 文件不存在: {source}")

        from docx import Document as DocxDocument

        doc = DocxDocument(source)

        # 1. 提取段落文本
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # 2. 多线程提取表格内容
        table_texts = self._extract_tables_parallel(doc)

        # 3. 合并并分片
        full_text = "\n".join(paragraphs + table_texts)
        full_text = self.clean_text(full_text)

        chunks = self._split_text(full_text, chunk_size, chunk_overlap)

        documents = []
        for idx, chunk in enumerate(chunks):
            documents.append(Document(
                text=chunk,
                metadata={
                    "source": source,
                    "file_name": os.path.basename(source),
                    "chunk_index": idx,
                    "file_type": "word",
                },
            ))
        return documents

    def _extract_tables_parallel(self, doc) -> List[str]:
        """多线程提取表格文本。"""
        if not doc.tables:
            return []

        def extract_table(table) -> str:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            return "\n".join(rows)

        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            results = list(executor.map(extract_table, doc.tables))
        return results

    def _split_text(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        if not text:
            return []
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start = end - chunk_overlap
        return chunks
